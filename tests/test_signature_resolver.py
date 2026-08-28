"""Tests for app/food_cell/signature_resolver.py."""

from __future__ import annotations

import base64
import io
import os
from pathlib import Path
from unittest.mock import patch

import pytest


def _make_valid_jpeg() -> bytes:
    """Create a minimal valid JPEG that python-docx can parse."""
    try:
        from PIL import Image

        img = Image.new("RGB", (100, 30), color=(255, 255, 255))
        buf = io.BytesIO()
        img.save(buf, format="JPEG")
        return buf.getvalue()
    except ImportError:
        # Fallback: minimal JFIF header + tiny white image
        # This is a known-minimal valid JPEG (1x1 white pixel)
        return base64.b64decode(
            "/9j/4AAQSkZJRgABAQEASABIAAD/2wBDAAgGBgcGBQgHBwcJCQgKDBQNDAsL"
            "DBkSEw8UHRofHh0aHBwgJC4nICIsIxwcKDcpLDAxNDQ0Hyc5PTgyPC4zNDL/"
            "2wBDAQkJCQwLDBgNDRgyIRwhMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIyMjIy"
            "MjIyMjIyMjIyMjIyMjIyMjIyMjIyMjL/wAARCAAyACgDASIAAhEBAxEB/8QA"
            "HwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAFRABAQAAAAAAAAAAAAAAAAAAAAf/"
            "xAABEAICAQIEBQMFBQUAAAAAAQIDBBEABRIhBhMxQVEiFAcYEyIjMkFVCBx"
            "FDBGBkLBUnLR0fEjMEBSCRUqJzJXUWFiQ5OTo/RkdISExPYHcqKkxNTk9B"
            "QlJVU1RUWltcYXJ0ZWZnaGlqaWxtbm9jXV5fsICAwUGBgoKSlJUVFhmZmZm"
            "am5ydHV2d3h5eoOEhYaHiImKi4yNjo+QkZKTlJWWl5iZmpucnZ6foKGio6Sl"
            "pqeoqaqrrK2ur7CxsrO0tba3uLm6u7y9vr/AwcLDxMXGx8jJysvMzc7P0NHS"
            "09TV1tfa2drb3N3e3+Dh4uPk5ebn6Onq6+zt7u/w8fLz9PX29/j5+vv8/f7/"
            "AABEIAAEAAQMBgMBAAAAAAAAAAAAEQECAwAhEhNBUSFRYXGBkSIyQpH/2gAMAwEAAhEDEQA/A"
            "K//Z"
        )


# ---------------------------------------------------------------------------
# Test fixtures
# ---------------------------------------------------------------------------

@pytest.fixture()
def sig_dir(tmp_path):
    """Create a temporary signature directory with a valid JPEG."""
    d = tmp_path / "signature"
    d.mkdir()
    (d / "sumansaha.jpg").write_bytes(_make_valid_jpeg())
    return d


# ---------------------------------------------------------------------------
# get_signature_path
# ---------------------------------------------------------------------------

class TestGetSignaturePath:
    """get_signature_path() maps FSO name -> file path."""

    def test_returns_none_for_none_name(self):
        from app.food_cell.signature_resolver import get_signature_path
        assert get_signature_path(None) is None

    def test_returns_none_for_empty_name(self):
        from app.food_cell.signature_resolver import get_signature_path
        assert get_signature_path("") is None

    def test_returns_path_when_file_exists(self, sig_dir):
        from app.food_cell.signature_resolver import get_signature_path
        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            result = get_signature_path("Suman Saha")
            assert result is not None
            assert result.name == "sumansaha.jpg"
            assert result.is_file()

    def test_returns_none_when_file_missing(self, sig_dir):
        from app.food_cell.signature_resolver import get_signature_path
        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            assert get_signature_path("Unknown Officer") is None

    def test_name_lowered_and_spaces_removed(self, sig_dir):
        from app.food_cell.signature_resolver import get_signature_path
        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            result = get_signature_path("Suman Saha")
            assert result is not None
            assert result.name == "sumansaha.jpg"


# ---------------------------------------------------------------------------
# get_signature_data_uri
# ---------------------------------------------------------------------------

class TestGetSignatureDataUri:
    """get_signature_data_uri() returns base64-encoded data URI."""

    def test_returns_none_when_no_file(self, sig_dir):
        from app.food_cell.signature_resolver import get_signature_data_uri
        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            assert get_signature_data_uri("Unknown Officer") is None

    def test_returns_data_uri_when_file_exists(self, sig_dir):
        from app.food_cell.signature_resolver import get_signature_data_uri
        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            uri = get_signature_data_uri("Suman Saha")
            assert uri is not None
            assert uri.startswith("data:image/jpeg;base64,")
            payload = uri.split(",", 1)[1]
            decoded = base64.b64decode(payload)
            assert len(decoded) > 0


# ---------------------------------------------------------------------------
# get_signature_bytes
# ---------------------------------------------------------------------------

class TestGetSignatureBytes:
    """get_signature_bytes() returns raw bytes."""

    def test_returns_none_when_no_file(self, sig_dir):
        from app.food_cell.signature_resolver import get_signature_bytes
        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            assert get_signature_bytes("Unknown Officer") is None

    def test_returns_bytes_when_file_exists(self, sig_dir):
        from app.food_cell.signature_resolver import get_signature_bytes
        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            result = get_signature_bytes("Suman Saha")
            assert result is not None
            assert len(result) > 0


# ---------------------------------------------------------------------------
# Word converter integration
# ---------------------------------------------------------------------------

class TestSignatureInWordConverter:
    """Word converter auto-resolves signature when fso_name is provided."""

    def test_signature_embedded_in_docx(self, sig_dir):
        from docx import Document
        from app.food_cell.word_converter import ImprovementNoticeWordConverter

        ctx = {
            "notice_date": "2026-08-28",
            "fbo_name": "Test Restaurant",
            "fbo_address": "123 Park Street",
            "fbo_fssai": "LIC-001",
            "fso_name": "Suman Saha",
            "improvement_notice_ref": "IMP-001",
            "violations": [{"description": "Unclean", "section": "S.32"}],
            "actions": ["Clean premises"],
            "compliance_deadline": "2026-09-28",
            "enclosures": [],
        }

        with patch("app.food_cell.signature_resolver._SIGNATURE_DIR", sig_dir):
            docx_bytes = ImprovementNoticeWordConverter().build(ctx)

        doc = Document(io.BytesIO(docx_bytes))
        # The image should be embedded as a relationship
        rels = doc.part.rels
        image_rels = [r for r in rels.values() if "image" in r.reltype]
        assert len(image_rels) >= 1, "No image embedded in .docx"

    def test_no_signature_still_works(self):
        from docx import Document
        from app.food_cell.word_converter import ImprovementNoticeWordConverter

        ctx = {
            "notice_date": "2026-08-28",
            "fbo_name": "Test Restaurant",
            "fbo_address": "123 Park Street",
            "fbo_fssai": "LIC-001",
            "fso_name": "No Signature Officer",
            "improvement_notice_ref": "IMP-001",
            "violations": [{"description": "Unclean", "section": "S.32"}],
            "actions": ["Clean premises"],
            "compliance_deadline": "2026-09-28",
            "enclosures": [],
        }

        docx_bytes = ImprovementNoticeWordConverter().build(ctx)
        doc = Document(io.BytesIO(docx_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "No Signature Officer" in text
        assert "Food Safety Officer" in text


# ---------------------------------------------------------------------------
# HTML template integration
# ---------------------------------------------------------------------------

class TestSignatureInHTMLTemplate:
    """HTML template includes signature image when signature_path is provided."""

    def _render(self, signature_path=None, fso_name="Suman Saha"):
        from app import create_app
        from flask import render_template

        app = create_app()
        with app.app_context():
            return render_template(
                "food_cell/improvement_notice.html",
                signature_path=signature_path,
                fso_name=fso_name,
                notice_date="2026-08-28",
                fbo_name="Test",
                fbo_address="Addr",
                fbo_fssai="LIC-001",
                improvements=[],
                improvement_notice_ref="IMP-001",
                compliance_deadline="2026-09-28",
                violations=[],
                actions=[],
                enclosures=[],
            )

    def test_signature_img_in_rendered_html(self):
        sig_path = "/tmp/fake/sumansaha.jpg"
        html = self._render(signature_path=sig_path)
        assert f'src="file:///{sig_path}"' in html
        assert "Signature of Suman Saha" in html

    def test_no_signature_path_no_img(self):
        html = self._render(signature_path=None)
        assert "file:///" not in html
        assert "Suman Saha" in html
