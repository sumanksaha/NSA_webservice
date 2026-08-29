#!/usr/bin/env python3
"""
Build Word (.docx) documents for adjudication cases.

Uses python-docx for consistent Microsoft Word output matching the visual style
of the existing PDFs (WeasyPrint / Asciidoctor).
"""

from datetime import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


class AdjudicationWordConverter:
    """Utility to generate .docx files for Petition and Permission Letter."""

    # Shared formatting styles
    TITLE_STYLE = "Title"
    HEADER_STYLE = "Heading 1"
    BODY_STYLE = "Normal"
    FOOTER_STYLE = "Footer"

    def __init__(self):
        # Optional: load custom docx templates (pending deployment)
        pass

    # ---------------------------------------------------------------------
    # Public API
    # ---------------------------------------------------------------------

    def build_petition(self, context: dict) -> bytes:
        """Build a Word document from the AsciiDoc rendering context for Petition."""
        doc = Document()
        self._set_document_margins(doc)

        # Header section
        self._add_header_section(doc, context)
        self._add_address_block(doc, context)
        self._add_case_number(doc, context)
        self._add_body_section(doc, context)
        self._add_violations_section(doc, context)
        self._add_photo_evidence_section(doc, context)
        self._add_foot_notes(doc, context)
        self._add_encumbrances_sections(doc, context)

        # Add footer (compilation date)
        self._add_footer(doc, context)

        # Save to bytes stream
        stream = self._get_stream_bytes(doc)
        return stream

    def build_permission_letter(self, context: dict) -> bytes:
        """Build a Word document from the AsciiDoc rendering context for Permission Letter."""
        doc = Document()
        self._set_document_margins(doc)

        # Header
        self._add_letter_header(doc, context)
        self._add_letter_body(doc, context)
        self._add_letter_footer(doc, context)

        # Footer
        self._add_footer(doc, context)

        return self._get_stream_bytes(doc)

    # ---------------------------------------------------------------------
    # Internal helpers
    # ---------------------------------------------------------------------

    def _set_document_margins(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    def _add_header_section(self, doc: Document, context: dict) -> None:
        """Add the header with Ld. Adjudicating Officer, KMC and case info."""
        header = doc.add_paragraph()
        header.text = "Before the Ld. Adjudicating Officer, KMC"
        header.style = self.TITLE_STYLE

        case_no = context.get("case_number", "F.S. & S Case No.")
        para = doc.add_paragraph()
        para.text = f"F.S. & S Case No. {case_no}"
        para.style = self.BODY_STYLE
        doc.add_paragraph()  # spacing

    def _add_address_block(self, doc: Document, context: dict) -> None:
        """Two-column address block (Complainant / Licensee)."""
        # Left column: Food Safety Officer, License Officer
        left = doc.add_paragraph()
        left.text = "1. {food_safety_officer_name}\nFood Safety Officer\nKolkata Municipal Corporation"
        left.style = self.BODY_STYLE
        self._add_spacing_paragraph(doc)

        right = doc.add_paragraph()
        right.text = "2. License Officer\n5, S.N. Banerjee Road, Kolkata - 700013"
        right.style = self.BODY_STYLE
        self._add_spacing_paragraph(doc)

        doc.add_paragraph()  # big spacing

    def _add_case_number(self, doc: Document, context: dict) -> None:
        """Display the case number for the current Adjudication."""
        case_no = context.get("case_number", "F.S. & S Case No.")
        para = doc.add_paragraph()
        para.text = f"F.S. & S Case No. {case_no}"
        para.style = self.BODY_STYLE
        doc.add_paragraph()

    def _add_body_section(self, doc: Document, context: dict) -> None:
        """Add the Petition body (Most Respectfully Sheweth)."""
        para = doc.add_paragraph()
        para.text = "Most Respectfully Sheweth: -"
        para.style = self.BODY_STYLE
        para = doc.add_paragraph()
        para.text = "The humble petition on behalf of the complainant above named:"
        para.style = self.BODY_STYLE

        # Statement of Facts (we could potentially parse from context, but for now, a simple block)
        facts = self._format_facts(context)
        doc.add_paragraph(facts, style=self.BODY_STYLE)

    def _add_violations_section(self, doc: Document, context: dict) -> None:
        """Add the violation table if violations are present."""
        violations = context.get("violations", [])
        if not violations:
            return

        doc.add_paragraph("\nObserved Non-Compliances During Inspection:", style=self.HEADER_STYLE)
        # Create a simple table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Sl. No"
        header_cells[1].text = "Nature of Non-Compliance"
        header_cells[2].text = "Observation"

        for idx, violation in enumerate(violations, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = violation.get("title", "")
            row[2].text = violation.get("observation", "")

    def _add_photo_evidence_section(self, doc: Document, context: dict) -> None:
        """Add a table for photo evidence if any."""
        adjudication = context.get("adjudication", {})
        photos = adjudication.get("photos", [])
        if not photos:
            return

        doc.add_paragraph("\nPhoto Evidence:", style=self.HEADER_STYLE)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Sl. No"
        header_cells[1].text = "Image ID"
        header_cells[2].text = "Location"
        header_cells[3].text = "Captured At"

        for idx, photo in enumerate(photos, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = photo.get("id", "")
            row[2].text = photo.get("locality", "Unresolved")
            row[3].text = str(photo.get("captured_at", ""))

    def _add_foot_notes(self, doc: Document, context: dict) -> None:
        """Add foot notes for liability and acknowledgement."""
        applicable_sections = context.get("applicable_sections", [])
        if applicable_sections:
            doc.add_paragraph("\nLiability under Sections:", style=self.HEADER_STYLE)
            for sec in applicable_sections:
                doc.add_paragraph(f"Section {sec}: Liability provisions apply.", style=self.BODY_STYLE)

        # Foot notes and signature
        doc.add_paragraph(
            "\nUnder the circumstances, it is prayed that Your Honour may kindly be pleased to issue notice upon the FBO to make representation...",
            style=self.BODY_STYLE,
        )

        # Signature block (simplified)
        signature = doc.add_paragraph("\nFood Safety Officer", style=self.BODY_STYLE)
        signature.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Footer content (temporary)
        doc.add_paragraph(f"Compiled on: {datetime.today().strftime('%d %B %Y')}", style=self.FOOTER_STYLE)

    def _add_encumbrances_sections(self, doc: Document, context: dict) -> None:
        """Add the encumbrances sections for both petition and permission letter."""
        # If we had more specific context about encumbrances, we would add them here.
        # For now, we'll simply add a placeholder section.
        doc.add_paragraph("\nList of Enclosures:", style=self.HEADER_STYLE)
        doc.add_paragraph("1. Photocopies of Form V dated __________________", style=self.BODY_STYLE)
        doc.add_paragraph("2. Photocopies of Form VI dated __________________", style=self.BODY_STYLE)
        doc.add_paragraph("3. Lab submission receipt dated __________________", style=self.BODY_STYLE)
        doc.add_paragraph("4. Analyst Report __________________ dated __________________", style=self.BODY_STYLE)
        doc.add_paragraph("5. Directive __________________ dated __________________", style=self.BODY_STYLE)
        doc.add_paragraph("6. Authorization of Designated Officer dated __________________", style=self.BODY_STYLE)

    def _add_footer(self, doc: Document, context: dict) -> None:
        """Add a footer for the compilation date."""
        # The footer is typically added to all sections. Since we only have one section:
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = f"Compiled on: {datetime.today().strftime('%d %B %Y')}"
        paragraph.style = self.FOOTER_STYLE

    # ---------------------------------------------------------------------
    # Letter specific methods (Permission Letter)
    # ---------------------------------------------------------------------

    def _add_letter_header(self, doc: Document, context: dict) -> None:
        """Add the letter header for Permission Letter."""
        # To, The Designated Officer, Food Cell, Kolkata Municipal Corporation.
        to_para = doc.add_paragraph("To,", style=self.BODY_STYLE)
        to_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        to_para = doc.add_paragraph(
            "The Designated Officer, Food Cell, Kolkata Municipal Corporation", style=self.BODY_STYLE
        )
        to_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        doc.add_paragraph()  # spacing

        # Subject line
        subject_para = doc.add_paragraph(
            "*Subject:* Service of letter no. { directive_letter_no } dated { directive_letter_date } along with Food Analyst's report No. { analyst_report_no } dated { analyst_report_date } in connection with the sample of - { product_name } (batch No - { batch_no }, Date of Manufacturing - { mfg_date }, Date of expiry - { expiry_date }) sold by { retailer_fbo_name } ({ retailer_address }, Lic No - { retailer_fssai }) collected on { inspection_date } vide code No. { sample_code }",
            style=self.BODY_STYLE,
        )
        subject_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        doc.add_paragraph()  # spacing

    def _add_letter_body(self, doc: Document, context: dict) -> None:
        """Add the letter body text."""
        body_para = doc.add_paragraph(
            "In compliance to your order, an intimation letter vide no. {{ directive_letter_no }} dated {{ directive_letter_date }} along with Food Analyst's report No. {{ analyst_report_no }} dated {{ analyst_report_date }} declaring the sample of {{ product_name }} (batch No - {{ batch_no }}, Date of Manufacturing - {{ mfg_date }}, Date of expiry - {{ expiry_date }}) manufactured by {{ manufacturer_fbo_name }} ({{ manufacturer_address }}, Lic No - {{ manufacturer_fssai }}) and sold by {{ retailer_fbo_name }} ({{ retailer_address }}, Lic No - {{ retailer_fssai }}) collected on {{ inspection_date }} vide code No. {{ sample_code }} as {{ analysis_result }}, has been served upon to {{ retailer_name }} on {{ retailer_report_receive_date }} and {{ manufacturer_name }} on {{ manufacturer_report_receive_date }}."
        )
        body_para.style = self.BODY_STYLE
        body_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Add more lines as per the original template...

    def _add_letter_footer(self, doc: Document, context: dict) -> None:
        """Add the letter footer with signature and address."""
        # Signature line
        signature = doc.add_paragraph("{{ food_safety_officer_name }}", style=self.BODY_STYLE)
        signature.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph("Food Safety Officer", style=self.BODY_STYLE)
        doc.add_paragraph("Kolkata Municipal Corporation", style=self.BODY_STYLE)

    # ---------------------------------------------------------------------
    # Internal helpers
    # ---------------------------------------------------------------------

    def _format_facts(self, context: dict) -> str:
        """Format the facts from the context into a readable paragraph."""
        # For simplicity, we'll just join the relevant context fields
        facts = []
        if "facts" in context:
            facts.append(context["facts"])
        # Add more fields as needed
        return "\n".join(facts)

    def _add_spacing_paragraph(self, doc: Document) -> None:
        """Add a spacing paragraph."""
        doc.add_paragraph()

    def _get_stream_bytes(self, doc: Document) -> bytes:
        """Save the document to an in-memory stream and return its bytes."""
        import io

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream.getvalue()


if __name__ == "__main__":
    # Simple CLI test
    converter = AdjudicationWordConverter()
    print("Adjudication Word Converter loaded successfully.")
