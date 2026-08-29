#!/usr/bin/env python3
"""Word (.docx) document generator for case file petition and permission letter."""

import io

from docx import Document


class CaseFileWordConverter:
    """Converts case file data to Word (.docx) documents."""

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add a heading with proper formatting."""
        doc.add_heading(text, level=level)

    def _add_paragraph(self, doc: Document, text: str, style: str = "Normal"):
        """Add a paragraph with optional styling."""
        p = doc.add_paragraph(text, style=style)
        return p

    def _add_table_row(self, table, values: list):
        """Add a row to a table with the given values."""
        row = table.add_row()
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = str(val) if val is not None else ""

    def build_petition(self, ctx: dict) -> bytes:
        """Build the Petition document from case data context.

        Args:
            ctx: Dictionary containing case file data (processed form data)

        Returns:
            BytesIO buffer content as bytes for download
        """
        doc = Document()

        # Title
        self._add_heading(doc, "PETITION", level=0)

        # Case header info as table
        table = doc.add_table(rows=1, cols=2)
        self._add_table_row(table, ["Case No:", ctx.get("case_number", "")])
        self._add_table_row(table, ["Date:", ctx.get("authorization_date", "")])

        # Addressing section
        doc.add_paragraph(
            f"The humble petition on behalf of {ctx.get('manufacturer_name', '')} "
            f"\nrepresented by {ctx.get('manufacturer_fbo_name', '')}, a Food Business Operator "
            f"registered under FSSAI License No. {ctx.get('manufacturer_fssai', '')}, "
            f"\naforesaid {ctx.get('product_name', '')} (Batch No.: {ctx.get('batch_no', '')}), "
            "\nprays that this Hon'ble Authority would be pleased to:-"
        )

        self._add_heading(doc, "STATEMENT OF FACTS", level=2)
        self._add_paragraph(doc, ctx.get("facts", "No facts provided."), style="Normal")

        # Violations table
        violations = ctx.get("violations", [])
        applicable_sections = ctx.get("applicable_sections_display", "")
        if violations or applicable_sections:
            table = doc.add_table(rows=len(violations) + 1, cols=4)
            self._add_table_row(table, ["S. No.", "Regulation", "Clause", "Violation"])
            for idx, v in enumerate(violations, 1):
                self._add_table_row(
                    table, [str(idx), v.get("regulation", ""), v.get("clause", ""), v.get("violation", "")]
                )

        # Prayer
        self._add_heading(doc, "PRAYER", level=2)
        self._add_paragraph(
            doc,
            "In view of the above, it is most respectfully prayed that this Hon'ble Authority "
            "would be pleased to:\n"
            "1. Direct the respondent to comply with the said provisions of the Act.\n"
            "2. Pass such order as deemed fit in the circumstances.\n"
            "3. Award costs of proceedings to be paid by the respondent.",
            style="Normal",
        )

        # Buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def build_permission_letter(self, ctx: dict) -> bytes:
        """Build the Permission Letter document from case data context.

        Args:
            ctx: Dictionary containing case file data (processed form data)

        Returns:
            BytesIO buffer content as bytes for download
        """
        doc = Document()

        # Title
        self._add_heading(doc, "PERMISSION LETTER", level=0)

        # Date header
        doc.add_paragraph(f"Date: {ctx.get('authorization_date', '')}")
        doc.add_paragraph("")

        # To section
        self._add_paragraph(
            doc, f"To,\nThe Food Safety Officer,\n{ctx.get('food_safety_officer_name', 'Food Safety Officer')}"
        )
        doc.add_paragraph("")

        # Subject
        self._add_heading(doc, "Subject: Request for Sample Collection and Analysis", level=2)

        # Body
        self._add_paragraph(doc, "Respected Sir/Madam,")
        self._add_paragraph(doc, "")

        self._add_paragraph(
            doc,
            f"I/We, {ctx.get('manufacturer_name', '')}, holding FSSAI License No. "
            f"{ctx.get('manufacturer_fssai', '')}, represented by {ctx.get('manufacturer_fbo_name', '')}, "
            f"do hereby request you to collect and analyze the said {ctx.get('product_name', '')} "
            f"(Batch No.: {ctx.get('batch_no', '')}) kept/preserved at "
            f"{ctx.get('manufacturer_address', '')}.",
        )

        self._add_paragraph(doc, "")
        self._add_paragraph(doc, "Yours faithfully,")
        self._add_paragraph(doc, "")

        # Signature line
        table = doc.add_table(rows=2, cols=3)
        self._add_table_row(table, ["", "", f"{ctx.get('manufacturer_fbo_name', '')}"])
        self._add_table_row(table, ["", "", "Signature / Designation"])

        # Buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
