"""Word (.docx) converter for Improvement Notices (u/s 32, FSS Act).

Builds a .docx from the same Jinja2-style context dict as the HTML template,
using simple black-and-white government styling: plain text, bold headings,
black-bordered tables, no colors, badges, or shading (matching the
simplified HTML template).

Example:
    from app.food_cell.word_converter import ImprovementNoticeWordConverter

    converter = ImprovementNoticeWordConverter()
    docx_bytes = converter.build(context)
"""

from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT  # noqa: F401  (kept for parity with section setup)
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

logger = logging.getLogger(__name__)


class ImprovementNoticeWordConverter:
    """Build a .docx Improvement Notice from a Jinja2-style context dict.

    Expected keys (same as the HTML template):
        fbo_name, fbo_address, inspection_date, fbo_fssai, fso_name,
        notice_date, improvement_notice_ref, violations (list of dicts),
        actions (list of str), compliance_deadline, enclosures (list of str).
    """

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #

    def build(self, context: dict[str, Any]) -> bytes:
        """Return the .docx file as *bytes*."""
        # Auto-resolve signature image if not already provided
        if "_signature_bytes" not in context:
            from app.food_cell.signature_resolver import get_signature_bytes

            sig = get_signature_bytes(context.get("fso_name"))
            if sig:
                context = {**context, "_signature_bytes": sig}

        doc = self._create_document()
        self._add_letterhead(doc)
        self._add_doc_class_line(doc)
        self._add_notice_date(doc, context)
        self._add_recipient(doc)
        self._add_body(doc, context)
        self._add_footer(doc, context)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ------------------------------------------------------------------ #
    # Internal helpers
    # ------------------------------------------------------------------ #

    def _create_document(self) -> Any:
        doc = Document()
        # Narrow margins for an official letter
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        # Default font (black)
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)
        return doc

    # ── Letterhead ──────────────────────────────────────────────────────

    def _add_letterhead(self, doc: Document) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("KOLKATA MUNICIPAL CORPORATION")
        run.font.size = Pt(15)
        run.bold = True

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("Food Safety Department — Under Food Safety and Standards Act, 2006")
        r2.font.size = Pt(10)

        # Black bottom border
        self._add_bottom_border(doc, color="000000", width=12)

    # ── Document type line ──────────────────────────────────────────────

    def _add_doc_class_line(self, doc: Document) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(12)
        run = p.add_run("IMPROVEMENT NOTICE — SECTION 32, FSS ACT, 2006")
        run.font.size = Pt(10)
        run.bold = True

    # ── Notice date ─────────────────────────────────────────────────────

    def _add_notice_date(self, doc: Document, ctx: dict[str, Any]) -> None:
        p = doc.add_paragraph()
        p.space_before = Pt(6)
        run = p.add_run(f"Date: {ctx.get('notice_date') or '—'}")
        run.font.size = Pt(10)

    # ── Recipient ───────────────────────────────────────────────────────

    def _add_recipient(self, doc: Document) -> None:
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        lbl = p.add_run("To\n")
        lbl.bold = True

        p.add_run("The Designated Officer,\nFood Cell,\nKolkata Municipal Corporation.").font.size = Pt(10)

    # ── Body ────────────────────────────────────────────────────────────

    def _add_body(self, doc: Document, ctx: dict[str, Any]) -> None:
        # Reference
        ref = ctx.get("improvement_notice_ref")
        if ref:
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            run = p.add_run(f"Ref: {ref}")
            run.bold = True
            run.font.size = Pt(10)

        # Subject
        fbo = ctx.get("fbo_name") or "[FBO Name]"
        addr = ctx.get("fbo_address") or "[FBO Address]"
        date = ctx.get("inspection_date") or "[Inspection Date]"
        p = doc.add_paragraph()
        p.space_before = Pt(4)
        run = p.add_run(f"Subject: Inspection report regarding an inspection of {fbo} situated at {addr} on {date}.")
        run.bold = True
        run.font.size = Pt(10.5)

        # Salutation
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        p.add_run("Sir/Madam,").font.size = Pt(11)

        # ── FBO summary table ──────────────────────────────────────────
        if ctx.get("fbo_name") or ctx.get("fbo_fssai"):
            self._add_section_heading(doc, "DETAILS OF INSPECTION")
            rows_data = []
            if ctx.get("fbo_name"):
                rows_data.append(("FBO Name", ctx["fbo_name"]))
            if ctx.get("fbo_address"):
                rows_data.append(("Address", ctx["fbo_address"]))
            if ctx.get("fbo_fssai"):
                rows_data.append(("FSSAI License No.", ctx["fbo_fssai"]))
            if ctx.get("inspection_date"):
                rows_data.append(("Inspection Date", ctx["inspection_date"]))
            if ctx.get("fso_name"):
                rows_data.append(("Food Safety Officer", ctx["fso_name"]))
            if rows_data:
                self._add_kv_table(doc, rows_data)

        # ── Part 1: Inspection Findings ────────────────────────────────
        self._add_section_heading(doc, "PART 1 — INSPECTION FINDINGS")
        p = doc.add_paragraph()
        p.space_before = Pt(4)
        run = p.add_run("An inspection was performed at ")
        run.font.size = Pt(11)
        r_bold = p.add_run(f"{addr}")
        r_bold.bold = True
        r_bold.font.size = Pt(11)
        r_mid = p.add_run(" on ")
        r_mid.font.size = Pt(11)
        r_date = p.add_run(f"{date}")
        r_date.bold = True
        r_date.font.size = Pt(11)
        p.add_run(", and the following deviation was observed.").font.size = Pt(11)

        # Violations table
        violations = ctx.get("violations") or []
        if violations:
            self._add_violations_table(doc, violations)
        else:
            p = doc.add_paragraph()
            run = p.add_run("No specific deviations were recorded during the inspection.")
            run.italic = True

        # ── Part 2: Grounds ────────────────────────────────────────────
        self._add_section_heading(doc, "PART 2 — GROUNDS FOR IMPROVEMENT NOTICE")
        p = doc.add_paragraph()
        p.space_before = Pt(4)
        p.add_run(
            "Based on the following observations, an improvement notice "
            "u/s 32 may kindly be granted on the following ground:"
        ).font.size = Pt(11)

        actions = ctx.get("actions") or []
        if actions:
            self._add_actions_list(doc, actions)
        else:
            p = doc.add_paragraph()
            run = p.add_run("No specific remedial actions prescribed.")
            run.italic = True

        # ── Compliance deadline ─────────────────────────────────────────
        deadline = ctx.get("compliance_deadline")
        if deadline:
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            p.add_run(
                "The FBO is hereby directed to comply with the above observations "
                "and take the required corrective action on or before "
            ).font.size = Pt(10.5)
            r = p.add_run(f"{deadline}.")
            r.bold = True
            r.font.size = Pt(10.5)

        # ── Enclosures ─────────────────────────────────────────────────
        enclosures = ctx.get("enclosures") or []
        if enclosures:
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            r = p.add_run("Enclosures:")
            r.bold = True
            r.font.size = Pt(10)
            for enc in enclosures:
                p = doc.add_paragraph(style="List Number")
                p.add_run(enc).font.size = Pt(10)

        # ── Signature block ────────────────────────────────────────────
        self._add_signature_block(doc, ctx)

    # ── Tables ──────────────────────────────────────────────────────────

    def _add_kv_table(self, doc: Document, rows: list[tuple[str, str]]) -> None:
        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for i, (key, val) in enumerate(rows):
            cell_k = table.cell(i, 0)
            cell_k.width = Cm(4.5)
            p = cell_k.paragraphs[0]
            run = p.add_run(key)
            run.bold = True
            run.font.size = Pt(10)

            cell_v = table.cell(i, 1)
            p = cell_v.paragraphs[0]
            p.add_run(val).font.size = Pt(10)

        self._style_table_borders(table)

    def _add_violations_table(self, doc: Document, violations: list[dict[str, str]]) -> None:
        table = doc.add_table(rows=1 + len(violations), cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row (light gray, matching the HTML template's table headers)
        headers = ["Sl.", "Nature of Deviation", "Observation"]
        widths = [Cm(1.2), Cm(5.0), Cm(9.0)]
        for j, hdr in enumerate(headers):
            cell = table.cell(0, j)
            cell.width = widths[j]
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.bold = True
            run.font.size = Pt(9)
            self._shade_cell(cell, "EEEEEE")

        # Data rows
        for i, v in enumerate(violations):
            row_idx = i + 1
            # Sl.
            cell = table.cell(row_idx, 0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(str(i + 1)).font.size = Pt(10)

            # Title / Nature of Deviation
            cell = table.cell(row_idx, 1)
            p = cell.paragraphs[0]
            p.add_run(v.get("title") or v.get("description") or v.get("name", "")).font.size = Pt(10)

            # Observation / Section
            cell = table.cell(row_idx, 2)
            p = cell.paragraphs[0]
            p.add_run(v.get("observation") or v.get("section") or v.get("detail", "")).font.size = Pt(10)

        self._style_table_borders(table)

    def _add_actions_list(self, doc: Document, actions: list[str]) -> None:
        for i, action in enumerate(actions):
            p = doc.add_paragraph()
            p.space_before = Pt(2)
            p.space_after = Pt(4)
            p.add_run(f"{i + 1}.  ").font.size = Pt(11)
            p.add_run(action).font.size = Pt(11)

    # ── Section heading ─────────────────────────────────────────────────

    def _add_section_heading(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.space_before = Pt(14)
        p.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)

    # ── Signature block ─────────────────────────────────────────────────

    def _add_signature_block(self, doc: Document, ctx: dict[str, Any]) -> None:
        p = doc.add_paragraph()
        p.space_before = Pt(36)
        lbl = p.add_run("Issued by\n")
        lbl.font.size = Pt(9)

        # Thin black signature line
        self._add_bottom_border(doc, color="000000", width=4)

        # Embed signature image if available
        sig_bytes = ctx.get("_signature_bytes")
        if sig_bytes:
            sig_stream = io.BytesIO(sig_bytes)
            doc.add_picture(sig_stream, width=Inches(1.5))

        p2 = doc.add_paragraph()
        p2.space_before = Pt(4)
        r = p2.add_run(f"{ctx.get('fso_name') or '[FSO Name]'}\n")
        r.bold = True
        r.font.size = Pt(11)
        p2.add_run("Food Safety Officer\n").font.size = Pt(11)
        p2.add_run("Kolkata Municipal Corporation\n").font.size = Pt(11)
        p2.add_run(ctx.get("notice_date") or "—").font.size = Pt(11)

    # ── Footer ──────────────────────────────────────────────────────────

    def _add_footer(self, doc: Document, ctx: dict[str, Any]) -> None:
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run("This document is electronically generated.").font.size = Pt(8)
        ref = ctx.get("improvement_notice_ref")
        if ref:
            p.add_run(f"    Ref: {ref}").font.size = Pt(8)

    # ── Low-level formatting utilities ──────────────────────────────────

    @staticmethod
    def _shade_cell(cell: Any, fill_hex: str) -> None:
        """Apply background shading to a table cell."""
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr.append(shd)

    @staticmethod
    def _add_bottom_border(doc: Document, color: str = "000000", width: int = 8) -> None:
        """Add a thin bottom border via a new paragraph with bottom border."""
        p = doc.add_paragraph()
        p.space_before = Pt(2)
        p.space_after = Pt(2)
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(width))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    @staticmethod
    def _style_table_borders(table: Any, color: str = "000000") -> None:
        """Apply uniform black borders to a table."""
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        tbl_pr.append(borders)


__all__ = ["ImprovementNoticeWordConverter"]
